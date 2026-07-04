from __future__ import annotations

import json
import math
import random
import sys
import time
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "佛山本地智能气象出行服务系统需求分析文档.docx"
ASSET_DIR = ROOT / "generated_assets"
ASSET_DIR.mkdir(exist_ok=True)

FONT_SIMSUN = "SimSun"
FONT_SIMHEI = "SimHei"
FONT_KAI = "KaiTi"
BLUE = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)


def set_east_asia_font(run, font_name: str = FONT_SIMSUN):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font=FONT_SIMSUN, size=10.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.text = ""
    run = p.add_run(text)
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="999999", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def paragraph_border_bottom(paragraph, color="4F81BD", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_east_asia_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_east_asia_font(tail)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_SIMSUN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SIMSUN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 14, BLUE, 12, 8),
        ("Heading 3", 12, RGBColor(31, 78, 121), 8, 6),
    ]:
        style = styles[style_name]
        style.font.name = FONT_SIMHEI
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SIMHEI)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.3


def add_para(doc, text="", style=None, align=None, first_line=True, size=10.5, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    if not first_line:
        p.paragraph_format.first_line_indent = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return p


def add_heading(doc, level: int, text: str):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_caption(doc, text: str, kind: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4 if kind == "figure" else 8)
    p.paragraph_format.space_after = Pt(8 if kind == "figure" else 4)
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = BLACK
    return p


def add_source_note(doc, text):
    p = add_para(doc, text, first_line=False, size=9, color=MUTED)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.hanging_indent = Pt(10.5)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run("（1）" if False else "• ")
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    return p


def add_numbered(doc, items):
    for idx, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.hanging_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(f"{idx}. {text}")
        set_east_asia_font(run)
        run.font.size = Pt(10.5)


def add_table(doc, caption: str, headers, rows, widths_cm):
    add_caption(doc, caption, "table")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "E8EEF5")
        set_cell_width(cell, widths_cm[i])
        set_cell_text(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_SIMHEI, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_width(cells[i], widths_cm[i])
            set_cell_text(cells[i], str(val), align=WD_ALIGN_PARAGRAPH.CENTER if widths_cm[i] < 2.4 else WD_ALIGN_PARAGRAPH.LEFT, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def load_font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simfang.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_round_rect(draw, box, fill, outline="#334155", width=2, radius=12):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered_text(draw, box, text, font, fill="#0f172a", line_gap=6):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 16
    lines = []
    for part in text.split("\n"):
        line = ""
        for ch in part:
            trial = line + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    heights = [draw.textbbox((0, 0), ln, font=font)[3] - draw.textbbox((0, 0), ln, font=font)[1] for ln in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for ln, h in zip(lines, heights):
        w = draw.textbbox((0, 0), ln, font=font)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), ln, font=font, fill=fill)
        y += h + line_gap


def draw_arrow(draw, start, end, color="#475569", width=3):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 14
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=color)


def create_system_structure_png():
    out = ASSET_DIR / "system_function_structure.png"
    im = Image.new("RGB", (1500, 980), "white")
    d = ImageDraw.Draw(im)
    title_f = load_font(34, True)
    node_f = load_font(24, True)
    small_f = load_font(20, False)
    d.text((430, 30), "佛山本地智能气象出行服务系统功能结构", font=title_f, fill="#0f172a")
    layers = [
        ("微信小程序端", 80, 130, "#DBEAFE", ["首页天气", "出行规划", "预警中心", "个人中心"]),
        ("Vue3管理后台", 80, 400, "#DCFCE7", ["数据概览", "气象管理", "出行内容", "用户反馈", "系统管理"]),
        ("SpringBoot服务层", 610, 130, "#FEF3C7", ["统一API", "气象服务", "出行推荐", "高德地图", "微信认证", "权限日志"]),
        ("数据与外部服务", 610, 570, "#FCE7F3", ["MySQL业务库", "高德天气/路线/POI", "微信code2Session", "气象预警规则"]),
    ]
    boxes = {}
    for layer, x, y, color, items in layers:
        draw_round_rect(d, (x, y, x + 390, y + 80), color, "#2563eb" if "小程序" in layer else "#166534" if "后台" in layer else "#92400e" if "服务" in layer else "#9d174d", 3)
        draw_centered_text(d, (x, y, x + 390, y + 80), layer, node_f)
        for i, item in enumerate(items):
            bx = x + (i % 2) * 195
            by = y + 115 + (i // 2) * 90
            draw_round_rect(d, (bx, by, bx + 170, by + 60), "#ffffff", "#94a3b8", 2, 8)
            draw_centered_text(d, (bx, by, bx + 170, by + 60), item, small_f)
        boxes[layer] = (x, y, x + 390, y + 80)
    draw_arrow(d, (470, 170), (610, 170))
    draw_arrow(d, (470, 440), (610, 220))
    draw_arrow(d, (805, 485), (805, 570))
    draw_arrow(d, (1000, 610), (1000, 380))
    draw_round_rect(d, (1100, 220, 1420, 470), "#F8FAFC", "#64748b", 2, 10)
    d.text((1130, 245), "核心业务闭环", font=node_f, fill="#0f172a")
    for i, item in enumerate(["气象数据同步", "规则匹配与风险评估", "路线与POI融合", "预警/反馈闭环"]):
        draw_centered_text(d, (1130, 300 + i * 38, 1390, 332 + i * 38), item, small_f)
    im.save(out)
    return out


def create_travel_flow_png():
    out = ASSET_DIR / "travel_query_flow.png"
    im = Image.new("RGB", (1500, 780), "white")
    d = ImageDraw.Draw(im)
    title_f = load_font(34, True)
    node_f = load_font(22, True)
    small_f = load_font(19, False)
    d.text((470, 30), "用户智能出行查询业务流程", font=title_f, fill="#0f172a")
    steps = [
        ("打开小程序\n授权定位", "#DBEAFE"),
        ("选择区域/地点\n输入出发地目的地", "#DBEAFE"),
        ("调用后端API\n提交出行时间与方式", "#FEF3C7"),
        ("高德地理编码\n路线/POI查询", "#FCE7F3"),
        ("匹配气象数据\n预警与风险路段", "#FEF3C7"),
        ("生成出行报告\n返回建议与风险", "#DCFCE7"),
    ]
    x0, y0, w, h, gap = 70, 170, 200, 115, 36
    centers = []
    for i, (label, fill) in enumerate(steps):
        x = x0 + i * (w + gap)
        draw_round_rect(d, (x, y0, x + w, y0 + h), fill, "#475569", 3, 14)
        draw_centered_text(d, (x, y0, x + w, y0 + h), label, node_f)
        centers.append((x + w, y0 + h / 2))
        if i > 0:
            prev = x0 + (i - 1) * (w + gap) + w
            draw_arrow(d, (prev + 5, y0 + h / 2), (x - 8, y0 + h / 2))
    details = [
        ("输入", "districtCode、originAddress、destinationAddress、departureTime、modeCode"),
        ("处理", "天气概况、路线摘要、风险等级、风险路段、推荐地点"),
        ("输出", "小程序卡片化展示，可跳转地图导航，并保留反馈入口"),
    ]
    for i, (k, v) in enumerate(details):
        y = 390 + i * 92
        draw_round_rect(d, (180, y, 1320, y + 62), "#F8FAFC", "#CBD5E1", 2, 8)
        d.text((215, y + 16), k, font=node_f, fill="#1D4ED8")
        d.text((320, y + 18), v, font=small_f, fill="#334155")
    im.save(out)
    return out


def excalidraw_id():
    return f"{int(time.time()*1000):x}{random.randint(1000, 9999):x}"


def rect_el(x, y, w, h, bg, stroke="#1e293b"):
    return {
        "id": excalidraw_id(), "type": "rectangle", "x": x, "y": y, "width": w, "height": h,
        "angle": 0, "strokeColor": stroke, "backgroundColor": bg, "fillStyle": "solid",
        "strokeWidth": 2, "strokeStyle": "solid", "roughness": 1, "opacity": 100,
        "groupIds": [], "frameId": None, "roundness": {"type": 3}, "seed": random.randint(1, 999999),
        "version": 1, "versionNonce": random.randint(1, 999999), "isDeleted": False,
        "boundElements": None, "updated": 1, "link": None, "locked": False,
    }


def text_el(x, y, text, size=20, color="#0f172a", w=180, h=40):
    return {
        "id": excalidraw_id(), "type": "text", "x": x, "y": y, "width": w, "height": h,
        "angle": 0, "strokeColor": color, "backgroundColor": "transparent", "fillStyle": "solid",
        "strokeWidth": 1, "strokeStyle": "solid", "roughness": 1, "opacity": 100,
        "groupIds": [], "frameId": None, "roundness": None, "seed": random.randint(1, 999999),
        "version": 1, "versionNonce": random.randint(1, 999999), "isDeleted": False,
        "boundElements": None, "updated": 1, "link": None, "locked": False,
        "text": text, "fontSize": size, "fontFamily": 5, "textAlign": "center",
        "verticalAlign": "middle", "containerId": None, "originalText": text, "lineHeight": 1.25,
    }


def arrow_el(x1, y1, x2, y2, label=None):
    el = {
        "id": excalidraw_id(), "type": "arrow", "x": x1, "y": y1, "width": x2 - x1, "height": y2 - y1,
        "angle": 0, "strokeColor": "#475569", "backgroundColor": "transparent", "fillStyle": "solid",
        "strokeWidth": 2, "strokeStyle": "solid", "roughness": 1, "opacity": 100,
        "groupIds": [], "frameId": None, "roundness": {"type": 2}, "seed": random.randint(1, 999999),
        "version": 1, "versionNonce": random.randint(1, 999999), "isDeleted": False,
        "boundElements": None, "updated": 1, "link": None, "locked": False,
        "points": [[0, 0], [x2 - x1, y2 - y1]], "lastCommittedPoint": None,
        "startBinding": None, "endBinding": None, "startArrowhead": None, "endArrowhead": "arrow",
    }
    return el


def save_excalidraw(path: Path, nodes, arrows, title):
    elements = [text_el(170, 20, title, 26, "#0f172a", 760, 50)]
    for x, y, w, h, label, fill in nodes:
        elements.append(rect_el(x, y, w, h, fill))
        elements.append(text_el(x + 10, y + 12, label, 18, "#0f172a", w - 20, h - 20))
    for a in arrows:
        elements.append(arrow_el(*a))
    data = {
        "type": "excalidraw",
        "version": 2,
        "source": "https://excalidraw.com",
        "elements": elements,
        "appState": {"viewBackgroundColor": "#ffffff", "gridSize": 20},
        "files": {},
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def create_excalidraw_files():
    save_excalidraw(
        ASSET_DIR / "system-function-structure.excalidraw",
        [
            (80, 110, 180, 80, "微信小程序端\n天气/出行/预警/我的", "#DBEAFE"),
            (80, 260, 180, 80, "Vue3管理后台\n运营管理与数据维护", "#DCFCE7"),
            (360, 180, 220, 100, "SpringBoot服务层\nAPI/认证/规则引擎", "#FEF3C7"),
            (700, 120, 190, 80, "MySQL业务库\n气象/用户/反馈", "#FCE7F3"),
            (700, 270, 190, 80, "外部平台\n高德地图/微信接口", "#E0E7FF"),
        ],
        [(260, 150, 360, 210), (260, 300, 360, 250), (580, 220, 700, 160), (580, 240, 700, 310)],
        "系统功能结构图",
    )
    save_excalidraw(
        ASSET_DIR / "travel-query-flow.excalidraw",
        [
            (60, 130, 150, 70, "用户授权定位", "#DBEAFE"),
            (250, 130, 160, 70, "填写出行条件", "#DBEAFE"),
            (450, 130, 170, 70, "后端接收请求", "#FEF3C7"),
            (660, 130, 170, 70, "路线与天气匹配", "#FCE7F3"),
            (870, 130, 170, 70, "返回出行报告", "#DCFCE7"),
        ],
        [(210, 165, 250, 165), (410, 165, 450, 165), (620, 165, 660, 165), (830, 165, 870, 165)],
        "用户出行查询流程图",
    )


def add_cover(doc):
    for _ in range(2):
        add_para(doc, "", first_line=False)
    p = add_para(doc, "广东东软学院", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    p.paragraph_format.space_after = Pt(10)
    p = add_para(doc, "综合能力实训", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True)
    p.paragraph_format.space_after = Pt(48)
    title = "基于SpringBoot+Vue3+微信小程序的佛山本地智能气象出行服务系统设计与实现"
    p = add_para(doc, title, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=17, bold=True)
    p.paragraph_format.line_spacing = 1.5
    p = add_para(doc, "需求分析文档", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, color=BLUE)
    paragraph_border_bottom(p, "4F81BD", "8")
    add_para(doc, "", first_line=False)
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "808080", "8")
    rows = [
        ("学    院", "计算机学院"),
        ("专    业", "软件工程"),
        ("班    级", ""),
        ("学    号", "25215150314"),
        ("学生姓名", "巫泉彬"),
        ("完成日期", "2026年7月"),
    ]
    for r, (label, value) in enumerate(rows):
        set_cell_width(table.rows[r].cells[0], 4.0)
        set_cell_width(table.rows[r].cells[1], 8.5)
        shade_cell(table.rows[r].cells[0], "F2F4F7")
        set_cell_text(table.rows[r].cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_SIMHEI, size=12)
        set_cell_text(table.rows[r].cells[1], value, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()


def add_toc(doc):
    add_para(doc, "目  录", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    items = [
        "第1章 项目概述",
        "第2章 用户角色与业务场景分析",
        "第3章 功能性需求分析",
        "第4章 非功能性需求分析",
        "第5章 核心业务流程分析",
        "第6章 数据需求分析",
        "第7章 运行环境与外部接口需求",
        "第8章 可行性与风险分析",
        "参考文献",
    ]
    for item in items:
        p = add_para(doc, item, first_line=False, size=11)
        p.paragraph_format.left_indent = Cm(1.0)
    doc.add_page_break()


def add_reference_basis(doc):
    rows = [
        ("服务区域", "佛山市下辖禅城、南海、顺德、高明、三水五区", "用于确定区域切换、区级气象数据和后台区域维护范围"),
        ("人口规模", "第七次全国人口普查公布佛山市常住人口约949.89万人", "说明气象出行服务具有较大的本地用户基础"),
        ("气候特征", "佛山属亚热带季风气候，夏季高温多雨，4-9月降水与强对流影响较集中", "支撑暴雨、高温、台风外围影响等场景化需求"),
        ("预警规则", "广东省气象灾害预警信号包含暴雨、台风、高温、雷雨大风等类型，并按颜色分级", "用于设计预警展示、提醒阈值和防御指引"),
        ("外部接口", "高德开放平台天气、地理编码、路线规划和POI能力支持按城市/区县编码查询", "支撑系统的天气同步、路线规划和地点推荐"),
        ("用户认证", "微信小程序code2Session机制可通过临时登录凭证换取openid", "支撑小程序用户登录、收藏地点和反馈提交"),
    ]
    add_table(doc, "表1.1 需求分析采用的数据与规范依据", ["依据项", "数据或规范内容", "对本系统需求的影响"], rows, [3.0, 6.0, 7.2])
    add_source_note(doc, "数据来源：佛山市第七次全国人口普查公报、广东省气象灾害预警信号发布规定、高德开放平台Web服务API文档、微信小程序官方接口文档等公开资料整理。")


def build_document(output_path: Path | None = None):
    create_excalidraw_files()
    structure_png = create_system_structure_png()
    flow_png = create_travel_flow_png()
    output_path = output_path or OUT_DOCX

    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_toc(doc)

    # Footer after cover/TOC
    for sec in doc.sections:
        add_page_number(sec.footer.paragraphs[0])

    add_heading(doc, 1, "第1章 项目概述")
    add_heading(doc, 2, "1.1 项目背景")
    add_para(doc, "佛山地处珠三角腹地，城市通勤、产业园区出行、短途旅游和商圈消费活动频繁。受亚热带季风气候影响，佛山夏季高温多雨，汛期暴雨、雷雨大风、台风外围环流和持续高温等天气过程容易对市民出行造成影响。传统天气查询工具通常以城市级天气展示为主，缺少对佛山五区、镇街、通勤路线、易积水路段和本地景点的场景化关联。")
    add_para(doc, "本项目拟建设“佛山本地智能气象出行服务系统”，以微信小程序作为普通用户入口，以Vue3管理后台作为运营维护入口，以SpringBoot后端作为统一服务支撑，将气象数据、地图路线、POI地点、风险路段和用户偏好进行整合，为佛山本地用户提供天气查询、预警提醒、智能出行建议和本地目的地推荐服务。")
    add_reference_basis(doc)

    add_heading(doc, 2, "1.2 项目目标")
    add_numbered(doc, [
        "面向佛山五区用户提供实时天气、短期预报、生活指数和气象预警信息，支持按区县和常用地点进行查询。",
        "结合天气状态、降水概率、风力、预警等级、路线方式和风险路段，生成可操作的出行建议。",
        "建设面向运营人员的Web管理后台，实现区域、预警、景点、风险路段、用户反馈和系统参数的维护。",
        "通过微信登录、收藏地点和意见反馈能力，形成用户侧使用闭环，提高系统的持续运营价值。",
    ])
    add_heading(doc, 2, "1.3 项目范围")
    rows = [
        ("服务区域", "佛山市禅城区、南海区、顺德区、高明区、三水区，后续可扩展到珠三角其他城市。"),
        ("用户终端", "微信小程序用户端、PC Web管理后台、SpringBoot后端API服务。"),
        ("核心功能", "气象查询、智能出行报告、预警中心、本地地点推荐、风险路段提示、用户收藏与反馈、后台运营管理。"),
        ("边界说明", "系统以气象辅助出行为核心，不承担实时公交调度、网约车派单、应急指挥调度等强实时业务。地图导航依托第三方地图平台。"),
    ]
    add_table(doc, "表1.2 项目范围说明", ["范围项", "说明"], rows, [3.2, 13.0])

    add_heading(doc, 1, "第2章 用户角色与业务场景分析")
    add_heading(doc, 2, "2.1 用户角色划分")
    rows = [
        ("普通用户", "佛山本地居民、通勤族、短途游客", "快速查询天气、了解预警、获得路线天气和风险提示。"),
        ("通勤用户", "有固定家/公司地点和出行时段的用户", "每日上班前获得通勤路线天气、降雨风险和出行方式建议。"),
        ("内容运营人员", "负责景点、风险路段、预警文案维护的后台用户", "维护出行地点、预警内容、生活指数和知识文档。"),
        ("系统管理员", "负责账号、权限、参数与日志的后台用户", "进行账号管理、系统参数配置、接口状态检查和操作审计。"),
    ]
    add_table(doc, "表2.1 用户角色与核心诉求", ["角色", "典型对象", "核心诉求"], rows, [3.0, 5.0, 8.2])
    add_heading(doc, 2, "2.2 典型业务场景")
    add_para(doc, "根据佛山本地气候和出行特点，系统重点覆盖日常通勤、短途出游、恶劣天气避险和后台运营四类场景。场景设计以“用户需要在较短时间内理解天气风险并作出出行决策”为核心。")
    add_numbered(doc, [
        "上班通勤场景：用户在早高峰前查看当前所在区天气、降雨概率和风险路段，系统建议是否改乘地铁、公交或延后出发。",
        "周末出游场景：用户根据天气选择千灯湖、西樵山、顺峰山等户外景点，或在降雨、高温时选择岭南天地、祖庙博物馆、商圈等室内目的地。",
        "预警避险场景：当暴雨、雷雨大风、高温等预警生效时，用户在小程序预警中心查看影响区域、防御指引和相关风险路段。",
        "后台运营场景：运营人员通过管理后台录入或修正景点、预警、风险路段、生活指数和用户反馈处理结果，使前端展示内容保持可维护。"
    ])

    add_heading(doc, 1, "第3章 功能性需求分析")
    add_para(doc, "系统功能采用“微信小程序端 + Vue3管理后台 + SpringBoot后端服务 + 数据与外部平台”的分层结构。功能结构如图3.1所示。")
    doc.add_picture(str(structure_png), width=Cm(15.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图3.1 系统功能结构图", "figure")
    add_source_note(doc, "图片来源：根据本系统代码结构和需求分析自制。")

    add_heading(doc, 2, "3.1 微信小程序端需求")
    add_heading(doc, 3, "3.1.1 气象信息服务模块")
    add_para(doc, "气象信息服务模块面向普通用户提供首页天气展示、区域切换、预报查询、预警中心和生活指数。系统应支持按佛山五区进行数据筛选，并在用户授权定位后优先展示所在区域数据。")
    rows = [
        ("WX-WE-01", "实时天气展示", "展示天气类型、温度、体感温度、湿度、风向风力、空气质量、降水概率、气压、能见度和更新时间。", "高"),
        ("WX-WE-02", "区域切换", "支持禅城、南海、顺德、高明、三水五区切换，并为后续镇街级扩展预留字段。", "高"),
        ("WX-WE-03", "天气预报", "展示多日预报，包括日期、星期、天气、最高/最低温、降水概率、风向风力和出行建议。", "高"),
        ("WX-WE-04", "预警中心", "展示生效中的预警类型、级别、影响区域、发布时间、失效时间和防御指引。", "高"),
        ("WX-WE-05", "生活指数", "展示穿衣、紫外线、洗车、运动、晾晒、旅游等指数及建议文案。", "中"),
    ]
    add_table(doc, "表3.1 小程序气象信息服务功能需求", ["编号", "功能名称", "需求描述", "优先级"], rows, [2.2, 3.2, 9.0, 1.8])

    add_heading(doc, 3, "3.1.2 智能出行服务模块")
    add_para(doc, "智能出行服务模块是本系统区别于普通天气应用的核心能力。用户输入出发地、目的地、出行时间和出行方式后，系统结合高德地图路线能力、当前天气、预报信息、风险路段和推荐地点生成出行报告。")
    rows = [
        ("WX-TR-01", "出行报告生成", "用户输入出发地、目的地、出行时间、方式后，返回天气概况、路线摘要、风险等级、出行建议。", "高"),
        ("WX-TR-02", "路线方式支持", "支持驾车、步行、骑行、公交/地铁等出行方式，后端按modeCode进行路线服务调用。", "高"),
        ("WX-TR-03", "风险路段提示", "当降雨、暴雨或台风等触发条件满足时，展示易积水、易拥堵或需绕行路段。", "高"),
        ("WX-TR-04", "目的地推荐", "按天气标签、室内/室外属性和推荐等级推荐佛山本地景点、商圈和文化场所。", "中"),
        ("WX-TR-05", "地图跳转", "展示地点坐标和路线概要，支持跳转第三方地图继续导航。", "中"),
    ]
    add_table(doc, "表3.2 小程序智能出行服务功能需求", ["编号", "功能名称", "需求描述", "优先级"], rows, [2.2, 3.2, 9.0, 1.8])

    add_heading(doc, 3, "3.1.3 用户中心模块")
    add_para(doc, "用户中心应围绕微信用户身份、常用地点和反馈闭环设计。系统通过微信临时登录凭证换取openid，并以JWT维护后续接口访问会话。")
    add_numbered(doc, [
        "微信登录：用户通过小程序登录接口完成身份识别，后端保存openid、昵称、头像和最近登录时间。",
        "收藏地点：用户可新增、查询、删除家、公司、常去景点等收藏地址，为通勤提醒和个性化推荐提供依据。",
        "意见反馈：用户可提交功能建议、气象纠错或体验问题，后台可标记处理状态并填写回复。",
        "个人资料：用户可维护昵称、头像和默认区域，系统不得在未授权情况下收集无关个人信息。",
    ])

    add_heading(doc, 2, "3.2 Web管理后台需求")
    add_para(doc, "管理后台基于Vue3、Vite、Element Plus、Pinia、Vue Router和ECharts实现，主要为运营人员和系统管理员提供数据查看、内容维护和系统配置能力。当前项目路由已包含数据大盘、气象管理、预警管理、出行管理、用户管理、反馈管理和系统管理等页面。")
    rows = [
        ("AD-01", "数据概览", "展示用户数、预警数量、出行地点、风险路段、反馈处理等统计指标。", "高"),
        ("AD-02", "区域管理", "维护佛山五区编码、名称、服务范围、经纬度、亮点和交通关注点。", "高"),
        ("AD-03", "预警管理", "新增、编辑、删除预警信息，维护预警类型、级别、影响区域和防御指引。", "高"),
        ("AD-04", "出行地点管理", "维护本地景点、商圈、室内外属性、天气标签、推荐等级和地址坐标。", "高"),
        ("AD-05", "风险路段管理", "维护易积水、易拥堵等风险路段及触发天气标签和绕行建议。", "高"),
        ("AD-06", "用户与反馈管理", "查看用户列表，处理用户反馈，维护状态和回复内容。", "中"),
        ("AD-07", "系统管理", "维护管理员账号、系统参数、生活指数和操作日志。", "中"),
    ]
    add_table(doc, "表3.3 Web管理后台功能需求", ["编号", "功能名称", "需求描述", "优先级"], rows, [2.0, 3.0, 9.5, 1.7])

    add_heading(doc, 2, "3.3 SpringBoot后端服务需求")
    add_para(doc, "后端负责统一API、数据持久化、气象同步、地图能力封装、微信登录、JWT认证和推荐逻辑。根据项目代码，后端主要包含DashboardController、TravelApiController、WechatUserController、AdminCrudController、AmapWeatherService、AmapRouteService、AmapPoiService、AmapGeocodingService、FoshanWeatherService和TravelPlanningService等组件。")
    rows = [
        ("BE-01", "气象数据服务", "提供当前天气、天气预报、预警、生活指数、知识文档和模块状态查询。", "/api/weather/current, /api/weather/forecast, /api/warnings/active"),
        ("BE-02", "智能出行服务", "根据区域、出发地、目的地、出行时间和方式生成出行报告。", "/api/travel/report"),
        ("BE-03", "地图服务封装", "封装高德地理编码、逆地理编码、驾车、步行、骑行、公交路线、POI搜索。", "/api/travel/geocode, /api/travel/route/*"),
        ("BE-04", "微信用户服务", "完成微信登录、资料更新、收藏地点和反馈提交。", "/api/wechat/login, /api/wechat/favorites, /api/wechat/feedback"),
        ("BE-05", "后台管理服务", "为后台提供区域、地点、风险路段、预警、用户、反馈、参数等CRUD接口。", "/api/admin/*"),
    ]
    add_table(doc, "表3.4 后端服务接口需求", ["编号", "服务名称", "需求描述", "关联接口"], rows, [2.0, 3.0, 7.0, 4.2])

    add_heading(doc, 1, "第4章 非功能性需求分析")
    add_heading(doc, 2, "4.1 性能需求")
    rows = [
        ("小程序首页加载", "首屏主要天气卡片在2秒内完成展示", "使用缓存数据、接口聚合和轻量化页面结构。"),
        ("气象查询接口", "常规查询响应时间不超过1秒", "后端按区域读取最近一条观测数据，减少复杂计算。"),
        ("出行报告接口", "常规情况下1.5秒内返回，第三方接口异常时3秒内降级返回基础建议", "路线服务调用需设置超时和降级策略。"),
        ("后台列表查询", "常规列表查询响应不超过2秒", "分页查询、索引优化和必要的筛选条件。"),
        ("数据同步频率", "实时天气每30分钟同步一次，预警信息应支持更高频或手动刷新", "当前WeatherSyncScheduler已按30分钟固定频率设计。"),
    ]
    add_table(doc, "表4.1 性能需求指标", ["指标项", "目标值", "实现约束"], rows, [3.5, 5.5, 7.2])
    add_heading(doc, 2, "4.2 可靠性需求")
    add_numbered(doc, [
        "第三方天气或地图接口不可用时，系统应返回本地数据库最近一次有效数据，并提示数据更新时间。",
        "预警数据、风险路段和用户反馈等核心业务数据应持久化到数据库，避免服务重启后丢失。",
        "后台管理操作应记录操作日志，重要配置变更应可追溯。",
        "数据库应定期备份，恢复策略至少覆盖最近7天的误删或异常更新场景。",
    ])
    add_heading(doc, 2, "4.3 安全性需求")
    add_numbered(doc, [
        "用户openid、收藏地点、定位区域和反馈内容属于个人信息或关联数据，系统应遵循最小必要原则进行采集和使用。",
        "后台接口必须经过JWT认证与角色权限校验，避免未授权访问预警发布、系统参数和管理员账号接口。",
        "微信AppSecret、高德API Key、JWT Secret等敏感配置不得硬编码到前端，应通过后端配置文件或环境变量管理。",
        "接口应对输入参数进行校验，避免SQL注入、越权删除、恶意提交反馈和重复请求等风险。",
    ])
    add_heading(doc, 2, "4.4 易用性与兼容性需求")
    add_para(doc, "小程序应保持“首页天气、出行、预警、我的”四个底部Tab，核心功能三步内可达；后台采用左侧菜单和模块化页面，列表、表单、筛选、编辑、删除等操作应保持一致。系统应兼容微信8.0及以上版本，后台兼容Chrome、Edge、Firefox近两个正式版本，并适配1920×1080及以上主流办公分辨率。")
    add_heading(doc, 2, "4.5 可扩展性需求")
    add_para(doc, "系统需为后续镇街级气象、积水实况、公交天气联动、同城活动推荐、向量知识库和更多城市扩展预留接口。后端服务应保持分层清晰，推荐规则、外部接口和业务实体应避免强耦合。")

    add_heading(doc, 1, "第5章 核心业务流程分析")
    add_heading(doc, 2, "5.1 用户出行查询流程")
    add_para(doc, "用户智能出行查询流程如图5.1所示。流程从小程序授权定位和填写出行条件开始，经过后端接口聚合、地图路线查询、天气与风险匹配，最终输出可读的出行报告。")
    doc.add_picture(str(flow_png), width=Cm(15.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图5.1 用户智能出行查询流程图", "figure")
    add_source_note(doc, "图片来源：根据TravelApiController和TravelPlanningService业务流程自制。")
    add_numbered(doc, [
        "用户打开小程序并授权定位，系统确定默认服务区域。",
        "用户填写出发地、目的地、出行时间和出行方式。",
        "小程序调用后端/api/travel/report接口提交查询参数。",
        "后端调用地理编码获取坐标，并按出行方式调用路线规划接口。",
        "后端读取当前天气、预警、风险路段和推荐地点，形成风险等级与出行建议。",
        "小程序展示天气摘要、路线概要、风险提示、推荐地点和地图跳转入口。",
    ])

    add_heading(doc, 2, "5.2 气象预警发布流程")
    add_numbered(doc, [
        "系统从第三方气象服务或后台人工录入获得预警信息。",
        "后台校验预警类型、级别、影响区域、发布时间、失效时间和防御指引。",
        "预警信息写入warning_notice表，并在状态为active时进入小程序预警中心。",
        "当微信订阅消息能力开通后，系统根据用户订阅范围与提醒阈值推送通知。",
        "预警失效后更新状态，前端不再作为生效预警展示，但后台保留历史记录。",
    ])
    add_heading(doc, 2, "5.3 后台内容维护流程")
    add_numbered(doc, [
        "管理员登录后台，系统验证账号密码并签发JWT。",
        "管理员进入对应模块，新增或编辑区域、出行地点、风险路段、生活指数、预警或系统参数。",
        "后端进行字段校验并写入数据库，同时记录操作日志。",
        "小程序端重新查询接口后获得最新数据，形成后台维护到用户展示的闭环。",
    ])

    add_heading(doc, 1, "第6章 数据需求分析")
    add_heading(doc, 2, "6.1 核心数据实体")
    rows = [
        ("District", "区域信息", "code、name、adminCode、latitude、longitude、serviceArea、highlights、transportFocus"),
        ("WeatherObservation", "实时气象", "district、observationTime、weatherType、temperature、apparentTemperature、humidity、windDirection、windScale、airQuality、precipitationProbability"),
        ("WeatherForecast", "天气预报", "district、forecastDate、weekLabel、weatherType、lowTemperature、highTemperature、precipitationProbability、travelAdvice"),
        ("WarningNotice", "预警信息", "warningType、severity、title、content、issuedAt、expiresAt、status、impactArea、defenseGuidance"),
        ("TravelPlace", "出行地点", "district、name、category、address、location、indoor、weatherTags、sceneTags、recommendLevel、highlight"),
        ("RiskSegment", "风险路段", "district、name、location、riskType、triggerWeatherTags、description、advice、priority"),
        ("User/UserFavoritePlace", "用户与收藏", "openid、nickname、avatarUrl、districtCode、favorite place name/address/tag"),
        ("Feedback", "意见反馈", "user、feedbackType、content、imageUrls、status、reply"),
        ("AdminUser/SystemParam/OperationLog", "后台管理", "管理员账号、角色、参数键值、操作内容、操作时间、IP"),
    ]
    add_table(doc, "表6.1 核心数据实体说明", ["实体", "含义", "主要字段"], rows, [3.5, 3.2, 9.5])
    add_heading(doc, 2, "6.2 数据来源与质量要求")
    add_numbered(doc, [
        "气象数据来源包括高德天气接口、气象公开数据和后台人工维护数据。系统应保存更新时间，避免用户误认为历史数据为实时数据。",
        "地图数据来源包括高德地理编码、逆地理编码、路线规划和POI搜索。系统应对失败结果进行提示或降级。",
        "用户数据来自微信登录授权和用户自主提交，收藏地点、反馈和区域偏好应支持用户删除或更新。",
        "运营数据由后台管理员维护，系统应通过表单校验减少无效坐标、重复地点和错误预警级别。",
    ])

    add_heading(doc, 1, "第7章 运行环境与外部接口需求")
    add_heading(doc, 2, "7.1 技术选型")
    rows = [
        ("后端", "Spring Boot 4.1.0、Java 21、Spring Web、Spring Data JPA、Spring Security、JJWT、MySQL Connector"),
        ("管理后台", "Vue 3、Vite、TypeScript、Element Plus、Pinia、Vue Router、Axios、ECharts"),
        ("微信小程序", "微信小程序原生框架，包含首页、出行、预警、我的四个Tab页面"),
        ("数据库", "MySQL 8.0，测试环境可使用H2；JPA自动维护实体表结构"),
        ("第三方接口", "高德天气、地理编码、路线规划、POI搜索；微信code2Session登录接口"),
    ]
    add_table(doc, "表7.1 系统技术选型与运行环境", ["类别", "说明"], rows, [3.0, 13.2])
    add_heading(doc, 2, "7.2 外部接口约束")
    add_para(doc, "高德天气接口以城市或区县adcode作为查询参数，本系统已配置佛山市adcode 440600以及五区adcode：禅城440604、南海440605、顺德440606、三水440607、高明440608。微信登录接口依赖小程序AppID、AppSecret和前端传入的临时code，后端应保护密钥安全并处理接口不可用时的异常。")

    add_heading(doc, 1, "第8章 可行性与风险分析")
    add_heading(doc, 2, "8.1 技术可行性")
    add_para(doc, "系统采用SpringBoot、Vue3和微信小程序原生框架，均为成熟技术栈。当前项目已包含后端实体、Repository、Controller、服务类、定时同步任务、后台页面和小程序页面，能够支撑需求分析中的主要功能。高德地图和微信登录能力均有官方接口文档，可通过标准HTTP接口集成。")
    add_heading(doc, 2, "8.2 经济与应用可行性")
    add_para(doc, "系统面向佛山市民日常出行和本地运营管理，不需要建设专用硬件。开发阶段可使用个人开发环境、MySQL数据库和第三方开放平台接口完成验证，部署阶段可采用普通云服务器和Nginx反向代理，成本可控。")
    add_heading(doc, 2, "8.3 风险与应对措施")
    rows = [
        ("第三方接口限额或失败", "天气、路线或POI查询不可用", "设置超时、错误提示、缓存最近数据，并保留后台手动维护能力。"),
        ("预警信息不及时", "用户可能错过恶劣天气提醒", "提高预警同步频率，后台支持人工发布和下架，记录发布时间。"),
        ("定位与隐私风险", "用户位置和收藏地址属于敏感个人信息", "按最小必要采集，提供授权提示和删除入口，接口统一鉴权。"),
        ("推荐规则过于简单", "复杂天气场景下建议不够准确", "将规则配置化，逐步引入历史反馈、风险路段和知识库优化。"),
    ]
    add_table(doc, "表8.1 项目风险与应对措施", ["风险项", "影响", "应对措施"], rows, [3.5, 4.2, 8.5])

    add_heading(doc, 1, "参考文献")
    refs = [
        "[1] 佛山市统计局. 佛山市第七次全国人口普查公报.",
        "[2] 广东省人民政府. 广东省气象灾害预警信号发布规定.",
        "[3] 高德开放平台. Web服务API：天气查询、地理编码、路径规划、POI搜索.",
        "[4] 微信开放文档. 小程序登录：code2Session接口说明.",
        "[5] 中华人民共和国全国人民代表大会常务委员会. 中华人民共和国个人信息保护法.",
        "[6] 项目源码：SpringBoot、Vue3、微信小程序端现有代码结构与接口实现.",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    path = build_document(target)
    print(path)
