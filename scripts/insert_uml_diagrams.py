from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path(r"D:\微软oneDriver\OneDrive\Desktop\weather")
INPUT_DOCX = WORKSPACE / "佛山本地智能气象出行服务系统需求分析文档.docx"
OUTPUT_DOCX = WORKSPACE / "佛山本地智能气象出行服务系统需求分析文档-UML插图版-v2.docx"
ASSET_DIR = WORKSPACE / "generated_assets" / "uml_diagrams"
PUML_DIR = ASSET_DIR / "puml"

WHITE = "#FFFFFF"
INK = "#1F2937"
MUTED = "#475569"
LIGHT_BORDER = "#94A3B8"
BLUE = "#DBEAFE"
BLUE_BORDER = "#2563EB"
GREEN = "#DCFCE7"
GREEN_BORDER = "#15803D"
YELLOW = "#FEF3C7"
YELLOW_BORDER = "#B45309"
PINK = "#FCE7F3"
PINK_BORDER = "#BE185D"
PURPLE = "#EDE9FE"
PURPLE_BORDER = "#7C3AED"
SLATE = "#F8FAFC"
GRAY_FILL = "#F3F4F6"
ORANGE = "#FFEDD5"
ORANGE_BORDER = "#C2410C"
RED = "#FEE2E2"
RED_BORDER = "#B91C1C"


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    PUML_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                r"C:\Windows\Fonts\msyhbd.ttc",
                r"C:\Windows\Fonts\simhei.ttf",
                r"C:\Windows\Fonts\arialbd.ttf",
            ]
        )
    candidates.extend(
        [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\arial.ttf",
        ]
    )
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


TITLE_FONT = font(50, bold=True)
SUBTITLE_FONT = font(32, bold=True)
LABEL_FONT = font(26, bold=True)
BODY_FONT = font(24)
SMALL_FONT = font(20)
CLASS_NAME_FONT = font(24, bold=True)
CLASS_ATTR_FONT = font(20)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, text_font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    for block in text.split("\n"):
        if not block:
            lines.append("")
            continue
        current = ""
        for ch in block:
            candidate = current + ch
            width = draw.textbbox((0, 0), candidate, font=text_font)[2]
            if width <= max_width or not current:
                current = candidate
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def draw_multiline_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    text_font: ImageFont.FreeTypeFont,
    fill: str = INK,
    line_spacing: int = 8,
) -> None:
    x0, y0, x1, y1 = box
    lines = wrap_text(draw, text, text_font, max(10, x1 - x0 - 24))
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line or "国", font=text_font)
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + max(0, len(lines) - 1) * line_spacing
    y = y0 + (y1 - y0 - total_h) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line or "国", font=text_font)
        w = bbox[2] - bbox[0]
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, font=text_font, fill=fill)
        y += h + line_spacing


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    text_font: ImageFont.FreeTypeFont = LABEL_FONT,
    radius: int = 28,
    width: int = 4,
    text_fill: str = INK,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_multiline_center(draw, box, text, text_font, fill=text_fill)


def ellipse_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    text_font: ImageFont.FreeTypeFont = BODY_FONT,
    width: int = 4,
) -> None:
    draw.ellipse(box, fill=fill, outline=outline, width=width)
    draw_multiline_center(draw, box, text, text_font)


def draw_title(draw: ImageDraw.ImageDraw, width: int, title: str, subtitle: str | None = None) -> None:
    bbox = draw.textbbox((0, 0), title, font=TITLE_FONT)
    title_w = bbox[2] - bbox[0]
    draw.text(((width - title_w) / 2, 36), title, font=TITLE_FONT, fill=INK)
    if subtitle:
        bbox = draw.textbbox((0, 0), subtitle, font=SMALL_FONT)
        sub_w = bbox[2] - bbox[0]
        draw.text(((width - sub_w) / 2, 102), subtitle, font=SMALL_FONT, fill=MUTED)


def dashed_line(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = LIGHT_BORDER,
    width: int = 2,
    dash: int = 12,
    gap: int = 8,
) -> None:
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if not length:
        return
    dx = (x2 - x1) / length
    dy = (y2 - y1) / length
    dist = 0.0
    while dist < length:
        seg = min(dash, length - dist)
        sx = x1 + dx * dist
        sy = y1 + dy * dist
        ex = x1 + dx * (dist + seg)
        ey = y1 + dy * (dist + seg)
        draw.line((sx, sy, ex, ey), fill=fill, width=width)
        dist += dash + gap


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    fill: str = "#334155",
    width: int = 3,
    dashed: bool = False,
    label: str | None = None,
    label_offset: int = -30,
) -> None:
    if dashed:
        dashed_line(draw, start, end, fill=fill, width=width)
    else:
        draw.line((start, end), fill=fill, width=width)

    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_len = 16
    arrow_w = 7
    p1 = (
        x2 - arrow_len * math.cos(angle) + arrow_w * math.sin(angle),
        y2 - arrow_len * math.sin(angle) - arrow_w * math.cos(angle),
    )
    p2 = (
        x2 - arrow_len * math.cos(angle) - arrow_w * math.sin(angle),
        y2 - arrow_len * math.sin(angle) + arrow_w * math.cos(angle),
    )
    draw.polygon([(x2, y2), p1, p2], fill=fill)

    if label:
        mx = (x1 + x2) / 2
        my = (y1 + y2) / 2 + label_offset
        lines = wrap_text(draw, label, SMALL_FONT, 280)
        widths = [draw.textbbox((0, 0), line, font=SMALL_FONT)[2] for line in lines]
        heights = [draw.textbbox((0, 0), line or "国", font=SMALL_FONT)[3] for line in lines]
        box_w = max(widths) + 24
        box_h = sum(heights) + max(0, len(lines) - 1) * 4 + 12
        label_box = (mx - box_w / 2, my - box_h / 2, mx + box_w / 2, my + box_h / 2)
        draw.rounded_rectangle(label_box, radius=12, fill=WHITE, outline="#E2E8F0", width=1)
        y = label_box[1] + 6
        for line, h, w in zip(lines, heights, widths):
            draw.text((mx - w / 2, y), line, font=SMALL_FONT, fill=INK)
            y += h + 4


def draw_actor(draw: ImageDraw.ImageDraw, center_x: int, top_y: int, label: str) -> None:
    head_r = 24
    draw.ellipse((center_x - head_r, top_y, center_x + head_r, top_y + head_r * 2), outline=INK, width=4)
    body_top = top_y + head_r * 2
    body_bottom = body_top + 70
    draw.line((center_x, body_top, center_x, body_bottom), fill=INK, width=4)
    draw.line((center_x - 35, body_top + 20, center_x + 35, body_top + 20), fill=INK, width=4)
    draw.line((center_x, body_bottom, center_x - 30, body_bottom + 45), fill=INK, width=4)
    draw.line((center_x, body_bottom, center_x + 30, body_bottom + 45), fill=INK, width=4)
    label_box = (center_x - 110, body_bottom + 54, center_x + 110, body_bottom + 120)
    draw_multiline_center(draw, label_box, label, LABEL_FONT)


def draw_external_system(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], name: str, fill: str, outline: str) -> None:
    rounded_box(draw, box, name, fill=fill, outline=outline, text_font=BODY_FONT)


def draw_participant(draw: ImageDraw.ImageDraw, center_x: int, name: str, fill: str, outline: str, top: int = 144) -> None:
    box = (center_x - 125, top, center_x + 125, top + 94)
    rounded_box(draw, box, name, fill=fill, outline=outline, text_font=SMALL_FONT, radius=18)
    dashed_line(draw, (center_x, top + 94), (center_x, 1420), fill="#CBD5E1", width=2)


def class_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    name: str,
    attrs: Iterable[str],
    methods: Iterable[str] | None = None,
    header_fill: str = BLUE,
    header_outline: str = BLUE_BORDER,
) -> None:
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=WHITE, outline=header_outline, width=3)
    header_h = 54
    draw.rounded_rectangle((x0, y0, x1, y0 + header_h), radius=18, fill=header_fill, outline=header_outline, width=3)
    draw.rectangle((x0, y0 + header_h - 18, x1, y0 + header_h), fill=header_fill, outline=header_fill)
    draw_multiline_center(draw, (x0 + 6, y0 + 6, x1 - 6, y0 + header_h - 4), name, CLASS_NAME_FONT)
    y = y0 + header_h + 12
    for attr in attrs:
        draw.text((x0 + 16, y), attr, font=CLASS_ATTR_FONT, fill=INK)
        y += 28
    if methods:
        draw.line((x0, y + 6, x1, y + 6), fill=LIGHT_BORDER, width=2)
        y += 18
        for method in methods:
            draw.text((x0 + 16, y), method, font=CLASS_ATTR_FONT, fill=MUTED)
            y += 28


def draw_multiplicity(draw: ImageDraw.ImageDraw, pos: tuple[int, int], text: str) -> None:
    bbox = draw.textbbox((0, 0), text, font=SMALL_FONT)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    x, y = pos
    draw.rounded_rectangle((x - 8, y - 6, x + w + 8, y + h + 6), radius=10, fill=WHITE, outline="#E2E8F0", width=1)
    draw.text((x, y), text, font=SMALL_FONT, fill=INK)


def build_use_case_diagram(out_path: Path) -> None:
    img = Image.new("RGB", (1800, 1300), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "系统核心用例图", "角色、后台运营与外部服务之间的职责边界")

    boundary = (360, 190, 1395, 1160)
    draw.rounded_rectangle(boundary, radius=28, outline="#334155", width=4, fill=SLATE)
    draw.text((boundary[0] + 24, boundary[1] + 20), "佛山本地智能气象出行服务系统", font=SUBTITLE_FONT, fill=INK)

    draw_actor(draw, 150, 270, "普通用户")
    draw_actor(draw, 150, 820, "系统管理员")
    draw_actor(draw, 1650, 260, "气象数据平台")
    draw_actor(draw, 1650, 520, "高德地图服务")
    draw_actor(draw, 1650, 780, "微信订阅消息")

    use_cases = {
        "查询实时天气": (500, 290, 840, 400, BLUE, BLUE_BORDER),
        "生成出行报告": (930, 300, 1270, 410, GREEN, GREEN_BORDER),
        "查看风险预警": (500, 470, 840, 580, YELLOW, YELLOW_BORDER),
        "管理常用地点": (930, 470, 1270, 580, PURPLE, PURPLE_BORDER),
        "维护天气数据": (500, 660, 840, 770, BLUE, BLUE_BORDER),
        "发布气象预警": (930, 660, 1270, 770, ORANGE, ORANGE_BORDER),
        "维护出行地点": (500, 850, 840, 960, GREEN, GREEN_BORDER),
        "处理用户反馈": (930, 850, 1270, 960, PINK, PINK_BORDER),
    }
    for name, (x0, y0, x1, y1, fill, outline) in use_cases.items():
        ellipse_box(draw, (x0, y0, x1, y1), name, fill, outline, text_font=BODY_FONT)

    for target in ["查询实时天气", "生成出行报告", "查看风险预警", "管理常用地点"]:
        x0, y0, x1, y1, *_ = use_cases[target]
        draw.line((240, (y0 + y1) // 2, x0 - 18, (y0 + y1) // 2), fill="#334155", width=3)
    for target in ["维护天气数据", "发布气象预警", "维护出行地点", "处理用户反馈"]:
        x0, y0, x1, y1, *_ = use_cases[target]
        draw.line((240, (y0 + y1) // 2, x0 - 18, (y0 + y1) // 2), fill="#334155", width=3)

    draw.line((1495, 355, 1278, 355), fill="#334155", width=3)
    draw.line((1495, 560, 1278, 355), fill="#334155", width=3)
    draw.line((1495, 835, 1278, 715), fill="#334155", width=3)

    img.save(out_path)


def build_travel_sequence_diagram(out_path: Path) -> None:
    img = Image.new("RGB", (2200, 1550), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "用户出行查询序列图", "对应 /api/travel/report 与 TravelPlanningService 核心交互")

    participants = [
        ("用户", BLUE, BLUE_BORDER),
        ("微信小程序", BLUE, BLUE_BORDER),
        ("TravelApi\nController", GREEN, GREEN_BORDER),
        ("TravelPlanning\nService", GREEN, GREEN_BORDER),
        ("FoshanWeather\nService", YELLOW, YELLOW_BORDER),
        ("AmapGeoRoute\nService", PURPLE, PURPLE_BORDER),
        ("高德开放平台", PINK, PINK_BORDER),
    ]
    xs = [170, 470, 790, 1110, 1430, 1750, 2050]
    for (name, fill, outline), x in zip(participants, xs):
        draw_participant(draw, x, name, fill, outline)

    messages = [
        (0, 1, 300, "1 输入出发地/目的地/时间"),
        (1, 2, 400, "2 GET /api/travel/report"),
        (2, 3, 520, "3 buildTravelReport(...)"),
        (3, 4, 650, "3.1 getTravelReport(...)"),
        (4, 3, 760, "3.1.1 返回天气摘要/风险等级", True),
        (3, 5, 900, "3.2 geocode(起点/终点)"),
        (5, 6, 1010, "3.2.1 地址解析"),
        (6, 5, 1110, "3.2.2 返回坐标", True),
        (3, 5, 1230, "3.3 route(origin,destination,mode)"),
        (5, 6, 1340, "3.3.1 路线规划"),
        (6, 5, 1440, "3.3.2 返回距离/时长/主指令", True),
        (3, 2, 1520, "4 组装 TravelReport", True),
        (2, 1, 1605, "5 返回天气摘要/路线/风险/建议", True),
        (1, 0, 1690, "6 展示出行报告并支持地图跳转", True),
    ]
    for item in messages:
        if len(item) == 4:
            from_idx, to_idx, y, label = item
            dashed = False
        else:
            from_idx, to_idx, y, label, dashed = item
        draw_arrow(draw, (xs[from_idx], y), (xs[to_idx], y), dashed=dashed, label=label)

    note_box = (980, 760, 1290, 860)
    draw.rounded_rectangle(note_box, radius=18, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw_multiline_center(
        draw,
        note_box,
        "TravelPlanningService\n综合天气快照、风险路段\n与推荐地点生成出行建议",
        SMALL_FONT,
    )

    img.save(out_path)


def build_warning_sequence_diagram(out_path: Path) -> None:
    img = Image.new("RGB", (2200, 1500), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "气象预警发布顺序图", "覆盖第三方预警接入、后台校验、状态入库与订阅消息推送")

    participants = [
        ("气象数据源/\n管理员", BLUE, BLUE_BORDER),
        ("管理后台", BLUE, BLUE_BORDER),
        ("WarningNotice\nService", GREEN, GREEN_BORDER),
        ("Warning\nRepository", YELLOW, YELLOW_BORDER),
        ("订阅筛选服务", PURPLE, PURPLE_BORDER),
        ("微信订阅消息", PINK, PINK_BORDER),
        ("用户小程序", ORANGE, ORANGE_BORDER),
    ]
    xs = [170, 470, 790, 1110, 1430, 1750, 2050]
    for (name, fill, outline), x in zip(participants, xs):
        draw_participant(draw, x, name, fill, outline)

    messages = [
        (0, 1, 300, "1 接收预警源数据或人工录入"),
        (1, 2, 420, "2 提交预警内容"),
        (2, 2, 530, "2.1 校验级别/区域/时效"),
        (2, 3, 660, "3 保存 warning_notice(status=active)"),
        (3, 2, 770, "3.1 返回入库结果", True),
        (2, 4, 900, "4 筛选订阅用户与提醒阈值"),
        (4, 5, 1030, "5 发送订阅消息"),
        (5, 6, 1160, "5.1 到达小程序通知入口"),
        (2, 6, 1290, "6 预警中心同步最新有效预警", True),
    ]
    for item in messages:
        if len(item) == 4:
            from_idx, to_idx, y, label = item
            dashed = False
        else:
            from_idx, to_idx, y, label, dashed = item
        if from_idx == to_idx:
            x = xs[from_idx]
            draw.line((x, y, x + 90, y), fill="#334155", width=3)
            draw.line((x + 90, y, x + 90, y + 60), fill="#334155", width=3)
            draw_arrow(draw, (x + 90, y + 60), (x, y + 60), label=label, label_offset=-24)
        else:
            draw_arrow(draw, (xs[from_idx], y), (xs[to_idx], y), dashed=dashed, label=label)

    note_box = (1320, 860, 1630, 940)
    draw.rounded_rectangle(note_box, radius=18, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw_multiline_center(draw, note_box, "订阅筛选规则\n= 区域范围 + 提醒阈值 + 订阅状态", SMALL_FONT)

    img.save(out_path)


def build_collaboration_diagram(out_path: Path) -> None:
    img = Image.new("RGB", (1900, 1350), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "后台内容维护协作图", "强调对象关系与编号消息，而不是时间轴")

    draw_actor(draw, 150, 380, "管理员")
    nodes = {
        "后台管理界面\n<<boundary>>": ((390, 300, 710, 440), BLUE, BLUE_BORDER),
        "AdminAuthController\n<<control>>": ((810, 210, 1150, 350), GREEN, GREEN_BORDER),
        "JwtUtil\n<<utility>>": ((1290, 230, 1600, 350), YELLOW, YELLOW_BORDER),
        "AdminCrudController\n<<control>>": ((790, 600, 1165, 750), GREEN, GREEN_BORDER),
        "WarningNoticeRepository\n<<entity>>": ((430, 980, 820, 1110), PURPLE, PURPLE_BORDER),
        "TravelPlaceRepository\n<<entity>>": ((890, 980, 1240, 1110), ORANGE, ORANGE_BORDER),
        "OperationLogRepository\n<<entity>>": ((1290, 980, 1650, 1110), PINK, PINK_BORDER),
    }
    for label, (box, fill, outline) in nodes.items():
        rounded_box(draw, box, label, fill, outline, text_font=BODY_FONT, radius=24)

    admin_anchor = (250, 490)
    draw_arrow(draw, admin_anchor, (390, 370), label="1 login()/选模块", label_offset=-78)
    draw_arrow(draw, (710, 330), (810, 275), label="1.1 auth()")
    draw_arrow(draw, (1150, 275), (1290, 275), label="1.2 issueToken()")
    draw_arrow(draw, (710, 395), (790, 675), label="2 submitCRUD()")
    draw_arrow(draw, (900, 750), (650, 980), label="2.1 saveWarning()", label_offset=-12)
    draw_arrow(draw, (1020, 750), (1060, 980), label="2.2 savePlace()")
    draw_arrow(draw, (1165, 675), (1290, 1045), label="2.3 recordLog()")
    draw_arrow(draw, (820, 1045), (790, 710), dashed=True, label="3.1 仓储返回", label_offset=44)
    draw_arrow(draw, (1240, 1045), (1165, 710), dashed=True, label="3.2 持久化返回")
    draw_arrow(draw, (790, 675), (710, 410), dashed=True, label="4 returnState()")
    draw_arrow(draw, (390, 405), admin_anchor, dashed=True, label="5 refresh()", label_offset=56)

    img.save(out_path)


def build_collaboration_diagram_standard(out_path: Path) -> None:
    img = Image.new("RGB", (2100, 1280), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "后台内容维护协作图", "UML Communication Diagram")

    def draw_comm_object(box: tuple[int, int, int, int], stereotype: str, class_name: str) -> None:
        x0, y0, x1, y1 = box
        draw.rectangle(box, fill=WHITE, outline=INK, width=3)

        stereotype_text = f"<<{stereotype}>>"
        name_text = f":{class_name}"

        stereo_bbox = draw.textbbox((0, 0), stereotype_text, font=SMALL_FONT)
        stereo_w = stereo_bbox[2] - stereo_bbox[0]
        stereo_h = stereo_bbox[3] - stereo_bbox[1]
        stereo_x = x0 + (x1 - x0 - stereo_w) / 2
        stereo_y = y0 + 18
        draw.text((stereo_x, stereo_y), stereotype_text, font=SMALL_FONT, fill=INK)

        name_bbox = draw.textbbox((0, 0), name_text, font=SMALL_FONT)
        name_w = name_bbox[2] - name_bbox[0]
        name_h = name_bbox[3] - name_bbox[1]
        name_x = x0 + (x1 - x0 - name_w) / 2
        name_y = stereo_y + stereo_h + 16
        draw.text((name_x, name_y), name_text, font=SMALL_FONT, fill=INK)
        draw.line((name_x, name_y + name_h + 4, name_x + name_w, name_y + name_h + 4), fill=INK, width=2)

    def draw_association(start: tuple[int, int], end: tuple[int, int]) -> None:
        draw.line((start, end), fill="#94A3B8", width=2)

    def draw_label(center: tuple[int, int], text: str, max_width: int = 430) -> None:
        cx, cy = center
        lines = wrap_text(draw, text, SMALL_FONT, max_width)
        widths = [draw.textbbox((0, 0), line or "A", font=SMALL_FONT)[2] for line in lines]
        heights = [draw.textbbox((0, 0), line or "A", font=SMALL_FONT)[3] for line in lines]
        box_w = max(widths) + 16
        box_h = sum(heights) + max(0, len(lines) - 1) * 4 + 12
        box = (cx - box_w / 2, cy - box_h / 2, cx + box_w / 2, cy + box_h / 2)
        draw.rectangle(box, fill=WHITE)
        y = box[1] + 6
        for line, h, w in zip(lines, heights, widths):
            draw.text((cx - w / 2, y), line, font=SMALL_FONT, fill=INK)
            y += h + 4

    def draw_message(
        start: tuple[int, int],
        end: tuple[int, int],
        label: str,
        label_center: tuple[int, int],
        *,
        dashed: bool = False,
    ) -> None:
        if dashed:
            dashed_line(draw, start, end, fill=INK, width=3, dash=12, gap=8)
        else:
            draw.line((start, end), fill=INK, width=3)

        x1, y1 = start
        x2, y2 = end
        angle = math.atan2(y2 - y1, x2 - x1)
        arrow_len = 18
        arrow_w = 8
        p1 = (
            x2 - arrow_len * math.cos(angle) + arrow_w * math.sin(angle),
            y2 - arrow_len * math.sin(angle) - arrow_w * math.cos(angle),
        )
        p2 = (
            x2 - arrow_len * math.cos(angle) - arrow_w * math.sin(angle),
            y2 - arrow_len * math.sin(angle) + arrow_w * math.cos(angle),
        )
        draw.line((x2, y2, p1[0], p1[1]), fill=INK, width=3)
        draw.line((x2, y2, p2[0], p2[1]), fill=INK, width=3)
        draw_label(label_center, label)

    draw_actor(draw, 120, 430, "管理员")

    boxes = {
        "page": (380, 470, 690, 590),
        "auth": (860, 240, 1210, 360),
        "jwt": (1390, 240, 1730, 360),
        "crud": (860, 730, 1210, 850),
        "warning": (1450, 540, 1930, 660),
        "travel": (1450, 750, 1930, 870),
        "log": (1450, 960, 1930, 1080),
    }
    draw_comm_object(boxes["page"], "boundary", "AdminPage")
    draw_comm_object(boxes["auth"], "control", "AdminAuthController")
    draw_comm_object(boxes["jwt"], "utility", "JwtUtil")
    draw_comm_object(boxes["crud"], "control", "AdminCrudController")
    draw_comm_object(boxes["warning"], "entity", "WarningNoticeRepository")
    draw_comm_object(boxes["travel"], "entity", "TravelPlaceRepository")
    draw_comm_object(boxes["log"], "entity", "OperationLogRepository")

    actor_page = ((170, 530), (380, 530))
    page_auth = ((690, 520), (860, 300))
    auth_jwt = ((1210, 300), (1390, 300))
    page_crud = ((690, 550), (860, 790))
    crud_warning = ((1210, 760), (1450, 600))
    crud_travel = ((1210, 790), (1450, 810))
    crud_log = ((1210, 820), (1450, 1010))

    for start, end in (actor_page, page_auth, auth_jwt, page_crud, crud_warning, crud_travel, crud_log):
        draw_association(start, end)

    draw_message(
        (170, 500),
        (380, 500),
        "1: requestLogin(username, password)",
        (275, 462),
    )
    draw_message(
        (170, 572),
        (380, 572),
        "2: submitMaintenance(operationType, payload)",
        (280, 618),
    )
    draw_message(
        (380, 536),
        (170, 536),
        "4: refreshedContentView",
        (270, 548),
        dashed=True,
    )
    draw_message(
        (690, 500),
        (860, 280),
        "1.1: authenticate(username, password)",
        (700, 392),
    )
    draw_message(
        (1210, 284),
        (1390, 284),
        "1.2: generateToken(adminClaims)",
        (1300, 242),
    )
    draw_message(
        (690, 562),
        (860, 802),
        "2.1: handleMaintenance(operationType, payload)",
        (820, 662),
    )
    draw_message(
        (1210, 746),
        (1450, 586),
        "2.1.1: saveWarningNotice(warningData)",
        (1316, 612),
    )
    draw_message(
        (1450, 624),
        (1210, 784),
        "2.1.2: warningNoticeResult",
        (1316, 726),
        dashed=True,
    )
    draw_message(
        (1210, 776),
        (1450, 796),
        "2.1.3: saveOrUpdateTravelPlace(placeData)",
        (1335, 742),
    )
    draw_message(
        (1450, 824),
        (1210, 804),
        "2.1.4: travelPlaceResult",
        (1325, 856),
        dashed=True,
    )
    draw_message(
        (1210, 836),
        (1450, 996),
        "2.1.5: saveOperationLog(logData)",
        (1322, 920),
    )
    draw_message(
        (1450, 1032),
        (1210, 872),
        "2.1.6: operationLogResult",
        (1320, 1038),
        dashed=True,
    )
    draw_message(
        (860, 818),
        (690, 578),
        "3: latestContentList, operationStatus",
        (650, 740),
        dashed=True,
    )

    img.save(out_path)


def build_class_diagram(out_path: Path) -> None:
    img = Image.new("RGB", (2200, 1500), WHITE)
    draw = ImageDraw.Draw(img)
    draw_title(draw, img.width, "系统核心类图", "基于 SpringBoot 实体类、服务类与主要对象关系抽取")

    rounded_box(draw, (850, 150, 1350, 270), "实体基类 BaseEntity\ncreatedAt / updatedAt / deleted", GRAY_FILL, LIGHT_BORDER, text_font=BODY_FONT, radius=20)
    class_box(draw, (80, 360, 460, 650), "User", ["- id: Long", "- openid: String", "- nickname: String", "- districtCode: String", "- status: Integer"], None, BLUE, BLUE_BORDER)
    class_box(draw, (80, 980, 460, 1230), "UserFavoritePlace", ["- id: Long", "- label: String", "- address: String", "- location: String"], None, BLUE, BLUE_BORDER)
    class_box(draw, (600, 520, 1040, 830), "District", ["- id: Long", "- code: String", "- name: String", "- serviceArea: String", "- latitude: Double", "- longitude: Double"], None, GREEN, GREEN_BORDER)
    class_box(draw, (1240, 240, 1740, 560), "WeatherObservation", ["- id: Long", "- observationTime: LocalDateTime", "- weatherType: String", "- temperature: BigDecimal", "- humidity: Integer"], None, YELLOW, YELLOW_BORDER)
    class_box(draw, (1240, 620, 1740, 930), "WarningNotice", ["- id: Long", "- warningType: String", "- severity: String", "- status: String", "- issuedAt: LocalDateTime"], None, PINK, PINK_BORDER)
    class_box(draw, (600, 1030, 1040, 1320), "TravelPlace", ["- id: Long", "- name: String", "- category: String", "- indoor: Boolean", "- recommendLevel: Integer"], None, ORANGE, ORANGE_BORDER)
    class_box(draw, (1240, 1030, 1740, 1320), "RiskSegment", ["- id: Long", "- name: String", "- riskType: String", "- triggerWeatherTags: String", "- priority: Integer"], None, PURPLE, PURPLE_BORDER)
    class_box(draw, (1800, 760, 2140, 1070), "TravelPlanningService", ["- foshanWeatherService", "- geocodingService", "- routeService"], ["+ buildTravelReport()"], GREEN, GREEN_BORDER)

    draw.line((270, 650, 270, 980), fill="#334155", width=3)
    draw_multiplicity(draw, (282, 760), "1")
    draw_multiplicity(draw, (282, 930), "0..*")

    draw.line((1040, 650, 1240, 420), fill="#334155", width=3)
    draw_multiplicity(draw, (1055, 600), "1")
    draw_multiplicity(draw, (1190, 450), "0..*")

    draw.line((1040, 700, 1240, 760), fill="#334155", width=3)
    draw_multiplicity(draw, (1060, 700), "1")
    draw_multiplicity(draw, (1190, 745), "0..*")

    draw.line((820, 830, 820, 1030), fill="#334155", width=3)
    draw_multiplicity(draw, (832, 880), "1")
    draw_multiplicity(draw, (832, 980), "0..*")

    draw.line((1040, 790, 1240, 1080), fill="#334155", width=3)
    draw_multiplicity(draw, (1060, 840), "1")
    draw_multiplicity(draw, (1200, 1040), "0..*")

    draw_arrow(draw, (1800, 820), (1740, 410), dashed=True, label="依赖天气摘要")
    draw_arrow(draw, (1800, 900), (1740, 760), dashed=True, label="依赖预警状态")
    draw_arrow(draw, (1800, 980), (1040, 1180), dashed=True, label="依赖推荐地点")
    draw_arrow(draw, (1800, 1040), (1740, 1180), dashed=True, label="依赖风险评估")

    img.save(out_path)


PUML_SOURCES = {
    "use_case": """@startuml
left to right direction
skinparam actorStyle awesome
skinparam packageStyle rectangle
actor "普通用户" as User
actor "系统管理员" as Admin
actor "气象数据平台" as WeatherProvider
actor "高德地图服务" as Amap
actor "微信订阅消息" as Wechat

rectangle "佛山本地智能气象出行服务系统" {
  usecase "查看实时天气" as UC1
  usecase "生成出行报告" as UC2
  usecase "查看预警与风险路段" as UC3
  usecase "管理常用地点与反馈" as UC4
  usecase "维护天气与区域数据" as UC5
  usecase "发布气象预警" as UC6
  usecase "维护出行地点/风险路段" as UC7
  usecase "处理反馈与系统参数" as UC8
}

User --> UC1
User --> UC2
User --> UC3
User --> UC4
Admin --> UC5
Admin --> UC6
Admin --> UC7
Admin --> UC8

UC1 ..> WeatherProvider : <<include>>
UC2 ..> WeatherProvider : <<include>>
UC2 ..> Amap : <<include>>
UC3 ..> WeatherProvider : <<include>>
UC6 ..> Wechat : <<include>>
@enduml
""",
    "travel_sequence": """@startuml
skinparam sequenceMessageAlign center
actor "用户" as User
participant "微信小程序" as MiniApp
participant "TravelApiController" as Api
participant "TravelPlanningService" as Planner
participant "FoshanWeatherService" as Weather
participant "AmapGeoRouteService" as AmapSvc
participant "高德开放平台" as Amap

User -> MiniApp : 输入出发地/目的地/时间
MiniApp -> Api : GET /api/travel/report
Api -> Planner : buildTravelReport(...)
Planner -> Weather : getTravelReport(...)
Weather --> Planner : 天气摘要/风险等级
Planner -> AmapSvc : geocode(起点/终点)
AmapSvc -> Amap : 地址解析
Amap --> AmapSvc : 返回坐标
Planner -> AmapSvc : route(origin,destination,mode)
AmapSvc -> Amap : 路线规划
Amap --> AmapSvc : 返回距离/时长/主指令
Planner --> Api : TravelReport
Api --> MiniApp : 天气摘要/路线/风险/建议
MiniApp --> User : 展示出行报告
@enduml
""",
    "warning_sequence": """@startuml
skinparam sequenceMessageAlign center
participant "气象数据源/管理员" as Source
participant "管理后台" as Backstage
participant "WarningNoticeService" as WarningService
database "WarningRepository" as Repo
participant "订阅筛选服务" as Filter
participant "微信订阅消息" as Wechat
participant "用户小程序" as MiniApp

Source -> Backstage : 提交预警信息
Backstage -> WarningService : 录入预警
WarningService -> WarningService : 校验级别/区域/时效
WarningService -> Repo : 保存 active 预警
Repo --> WarningService : 返回入库结果
WarningService -> Filter : 筛选订阅用户与阈值
Filter -> Wechat : 发送订阅消息
Wechat -> MiniApp : 到达通知入口
WarningService --> MiniApp : 同步最新有效预警
@enduml
""",
    "communication": """@startuml
left to right direction
skinparam shadowing false
skinparam defaultTextAlignment center
skinparam roundcorner 0
skinparam linetype polyline
actor "管理员" as Admin
boundary ":AdminPage" as Page
control ":AdminAuthController" as Auth
control ":AdminCrudController" as Crud
rectangle ":JwtUtil" <<utility>> as Jwt
entity ":WarningNoticeRepository" as WarningRepo
entity ":TravelPlaceRepository" as PlaceRepo
entity ":OperationLogRepository" as LogRepo

Admin -- Page
Page -- Auth
Auth -- Jwt
Page -- Crud
Crud -- WarningRepo
Crud -- PlaceRepo
Crud -- LogRepo

Admin -> Page : 1: login(username, password)
Page -> Auth : 1.1: authenticate(username, password)
Auth -> Jwt : 1.2: generateToken(adminClaims)
Admin -> Page : 2: submitContentOperation(operationType, payload)
Page -> Crud : 2.1: handleContentOperation(operationType, payload)
Crud -> WarningRepo : 2.1.1: saveWarningNotice(warningData)
WarningRepo --> Crud : 2.1.2: warningNoticeResult
Crud -> PlaceRepo : 2.1.3: saveOrUpdateTravelPlace(placeData)
PlaceRepo --> Crud : 2.1.4: travelPlaceResult
Crud -> LogRepo : 2.1.5: saveOperationLog(logData)
LogRepo --> Crud : 2.1.6: operationLogResult
Crud --> Page : 3: latestContentList, operationStatus
Page --> Admin : 4: refreshedContentView
@enduml
""",
    "class_diagram": """@startuml
skinparam classAttributeIconSize 0

abstract class BaseEntity {
  +createdAt
  +updatedAt
  +deleted
}

class User {
  -id: Long
  -openid: String
  -nickname: String
  -districtCode: String
  -status: Integer
}

class District {
  -id: Long
  -code: String
  -name: String
  -serviceArea: String
}

class WeatherObservation {
  -observationTime: LocalDateTime
  -weatherType: String
  -temperature: BigDecimal
  -humidity: Integer
}

class WarningNotice {
  -warningType: String
  -severity: String
  -status: String
  -issuedAt: LocalDateTime
}

class TravelPlace {
  -name: String
  -category: String
  -address: String
  -indoor: Boolean
  -recommendLevel: Integer
}

class RiskSegment {
  -name: String
  -riskType: String
  -triggerWeatherTags: String
  -priority: Integer
}

class UserFavoritePlace {
  -label: String
  -address: String
  -location: String
}

class TravelPlanningService {
  +buildTravelReport()
  +buildDrivingRoute()
}

BaseEntity <|-- User
BaseEntity <|-- District
BaseEntity <|-- WeatherObservation
BaseEntity <|-- WarningNotice
BaseEntity <|-- TravelPlace
BaseEntity <|-- RiskSegment
BaseEntity <|-- UserFavoritePlace

District "1" --> "0..*" WeatherObservation
District "1" --> "0..*" WarningNotice
District "1" --> "0..*" TravelPlace
District "1" --> "0..*" RiskSegment
User "1" --> "0..*" UserFavoritePlace

TravelPlanningService ..> WeatherObservation
TravelPlanningService ..> TravelPlace
TravelPlanningService ..> RiskSegment
@enduml
""",
}


def write_puml_sources() -> None:
    for name, content in PUML_SOURCES.items():
        (PUML_DIR / f"{name}.puml").write_text(content.strip() + "\n", encoding="utf-8")


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text is not None:
        new_para.add_run(text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str | None = None, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text is not None:
        new_para.add_run(text)
    return new_para


def style_caption(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)


def style_source(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)


def style_intro(paragraph: Paragraph, text: str) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)


def insert_image_before(paragraph: Paragraph, image_path: Path, width_in: float = 6.25) -> Paragraph:
    img_para = insert_paragraph_before(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(4)
    return img_para


def insert_image_block_after(
    paragraph: Paragraph,
    image_path: Path,
    intro: str,
    caption: str,
    source: str,
    width_in: float = 6.25,
) -> None:
    intro_para = insert_paragraph_after(paragraph)
    style_intro(intro_para, intro)
    img_para = insert_paragraph_after(intro_para)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(str(image_path), width=Inches(width_in))
    caption_para = insert_paragraph_after(img_para)
    style_caption(caption_para, caption)
    source_para = insert_paragraph_after(caption_para)
    style_source(source_para, source)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_diagrams(image_paths: dict[str, Path]) -> None:
    doc = Document(str(INPUT_DOCX))

    chapter2_anchor = find_paragraph(
        doc,
        "4. 后台运营场景：运营人员通过管理后台录入或修正景点、预警、风险路段、生活指数和用户反馈处理结果，使前端展示内容保持可维护。",
    )
    insert_image_block_after(
        chapter2_anchor,
        image_paths["use_case"],
        "系统主要参与者及其核心用例如图2.1所示，覆盖普通用户、后台运营以及外部服务协同。",
        "图2.1 系统核心用例图",
        "图片来源：根据第2章角色与业务场景分析、自有系统功能边界与外部服务关系绘制。",
    )

    fig51_caption = find_paragraph(doc, "图5.1 用户智能出行查询流程图")
    insert_image_before(fig51_caption, image_paths["travel_sequence"])
    style_caption(fig51_caption, "图5.1 用户出行查询序列图")
    fig51_source = find_paragraph(doc, "图片来源：根据TravelApiController和TravelPlanningService业务流程自制。")
    style_source(fig51_source, "图片来源：根据TravelApiController、TravelPlanningService与高德服务交互流程绘制。")

    chapter52_anchor = find_paragraph(
        doc,
        "5. 预警失效后更新状态，前端不再作为生效预警展示，但后台保留历史记录。",
    )
    insert_image_block_after(
        chapter52_anchor,
        image_paths["warning_sequence"],
        "气象预警发布、状态入库、订阅筛选与小程序同步顺序如图5.2所示。",
        "图5.2 气象预警发布顺序图",
        "图片来源：根据第5.2节预警发布流程与warning_notice数据流转逻辑绘制。",
    )

    chapter53_anchor = find_paragraph(
        doc,
        "4. 小程序端重新查询接口后获得最新数据，形成后台维护到用户展示的闭环。",
    )
    insert_image_block_after(
        chapter53_anchor,
        image_paths["communication"],
        "后台内容维护对象之间的职责分工与编号消息关系如图5.3所示。",
        "图5.3 后台内容维护协作图",
        "图片来源：根据AdminAuthController、AdminCrudController与OperationLog记录流程绘制。",
    )

    fig61_caption = find_paragraph(doc, "表6.1 核心数据实体说明")
    insert_image_before(fig61_caption, image_paths["class_diagram"], width_in=6.3)
    style_caption(fig61_caption, "图6.1 系统核心类图")
    fig61_source = insert_paragraph_after(fig61_caption)
    style_source(fig61_source, "图片来源：根据SpringBoot实体类、服务类与主要对象关系抽取绘制。")

    doc.save(str(OUTPUT_DOCX))


def main() -> None:
    ensure_dirs()
    write_puml_sources()

    images = {
        "use_case": ASSET_DIR / "use_case.png",
        "travel_sequence": ASSET_DIR / "travel_sequence.png",
        "warning_sequence": ASSET_DIR / "warning_sequence.png",
        "communication": ASSET_DIR / "communication.png",
        "class_diagram": ASSET_DIR / "class_diagram.png",
    }

    build_use_case_diagram(images["use_case"])
    build_travel_sequence_diagram(images["travel_sequence"])
    build_warning_sequence_diagram(images["warning_sequence"])
    build_collaboration_diagram_standard(images["communication"])
    build_class_diagram(images["class_diagram"])
    insert_diagrams(images)


if __name__ == "__main__":
    main()
